from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/research/agent-mesh-paper-v1.md"
OUTPUT = ROOT / "docs/research/agent-mesh-paper-v1.docx"

FONT = "Calibri"
INK = "17202A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT = "F4F6F9"


def set_font(run, size=None, bold=None, italic=None, color=INK, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeatable_field(paragraph, field):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, 8.5, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208
        style.paragraph_format.keep_together = True

    code = styles.add_style("Paper Code", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.2)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.line_spacing = 1.0

    abstract = styles.add_style("Paper Abstract", WD_STYLE_TYPE.PARAGRAPH)
    abstract.font.name = FONT
    abstract.font.size = Pt(10)
    abstract.paragraph_format.left_indent = Inches(0.25)
    abstract.paragraph_format.right_indent = Inches(0.25)
    abstract.paragraph_format.space_after = Pt(7)
    abstract.paragraph_format.line_spacing = 1.2
    abstract.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = FONT
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_together = True


def shade_paragraph(paragraph, fill=LIGHT):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def add_inline(paragraph, text, size=None):
    for part in INLINE.split(text):
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        italic = part.startswith("*") and part.endswith("*") and not bold
        code = part.startswith("`") and part.endswith("`")
        content = part[2:-2] if bold else part[1:-1] if italic or code else part
        run = paragraph.add_run(content)
        set_font(run, size=size, bold=bold or None, italic=italic or None,
                 color=INK, name="Courier New" if code else FONT)


def set_keep(paragraph):
    paragraph.paragraph_format.widow_control = True


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("Agent Mesh | Preprint v1.2"), 8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Page "), 8.5, color=MUTED)
    set_repeatable_field(footer, " PAGE ")

    # Publication-neutral first page: suitable for an anonymized preprint/review copy.
    title = lines[0].removeprefix("# ")
    subtitle = lines[2].removeprefix("## ")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(title), 23, bold=True, color="000000")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    set_font(p.add_run(subtitle), 13, italic=True, color=MUTED)

    in_code = False
    abstract_mode = False
    references_mode = False
    i = 3
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph(style="Paper Code")
            set_font(p.add_run(line), 9, color=INK, name="Courier New")
            set_keep(p)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        if line == "### Abstract":
            abstract_mode = True
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(7)
            set_font(p.add_run("Abstract"), 12, bold=True, color=DARK_BLUE)
            i += 1
            continue
        image_match = re.match(r"^!\[(.+)\]\((.+)\)$", line)
        if image_match:
            caption_text, relative_path = image_match.groups()
            image_path = (SOURCE.parent / relative_path).resolve()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            run = p.add_run()
            inline_shape = run.add_picture(str(image_path), width=Inches(6.45))
            inline_shape._inline.docPr.set("descr", caption_text)
            inline_shape._inline.docPr.set("title", caption_text.split(".", 1)[0])
            caption_p = doc.add_paragraph(style="Figure Caption")
            add_inline(caption_p, caption_text, size=9)
            set_keep(caption_p)
            i += 1
            continue
        if line.startswith("## "):
            abstract_mode = False
            references_mode = line == "## References"
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[3:])
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[4:])
        elif line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, line[5:])
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif line.startswith("**Douglas Rodriguez**") or line.startswith("douglas.rodriguez@") or line.startswith("Version 1."):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, line.rstrip(), size=10)
        else:
            style = "Paper Abstract" if abstract_mode else "Normal"
            p = doc.add_paragraph(style=style)
            add_inline(p, line, size=9.5 if references_mode else None)
            if abstract_mode:
                shade_paragraph(p)
            elif references_mode:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p.paragraph_format.space_after = Pt(7)
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.keep_together = True
        set_keep(p)
        i += 1

    doc.core_properties.title = "Agent Mesh: Governed Decentralized Coordination Under Partial Information"
    doc.core_properties.subject = "Architecture, control boundaries, and lessons from an inconclusive preregistered study"
    doc.core_properties.author = "Douglas Rodriguez"
    doc.core_properties.keywords = "multi-agent systems; decentralized coordination; agent governance; controlled emergence; DARPA DICE"
    doc.core_properties.comments = "Version 1.2 prepared for archival publication on Zenodo."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
